import os
import pandas as pd
import streamlit as st
import re
import copy
from PIL import Image

def get_salutation(name):
    name_upper = name.upper()
    
    # கம்பெனி பெயர்கள் (Company / Business)
    company_keywords = ["TRADERS", "MOTORS", "ENTERPRISES", "LIMITED", "LTD", "AGENCY", "WORKS", "STORES", "COMPANY", "CO"]
    if any(keyword in name_upper for keyword in company_keywords):
        return "M/S."
    
    # பெண்கள் பெயர்கள் (Female)
    female_keywords = ["MRS", "MISS", "KUMARI", "DEVI", "AMMAL", "MARY", "LAKSHMI", "ANITHA", "PRIYA", "KAVITHA"]
    if any(keyword in name_upper for keyword in female_keywords):
        return "MRS."
    
    # ஆண்கள் (Default Male)
    return "MR."

st.set_page_config(page_title="Auto Vehicle Quotation Generator", layout="wide")
st.title("🚗 Auto Vehicle Quotation Generator")

# --- 1. EXCEL DATA LOAD ---
@st.cache_data
def load_excel_data():
    file_path = "QA.xlsx"
    if os.path.exists(file_path):
        try:
            excel_file = pd.ExcelFile(file_path)
            sheet_name = "PRICE" if "PRICE" in excel_file.sheet_names else excel_file.sheet_names[0]
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            df.columns = df.columns.astype(str).str.strip()
            return df
        except Exception as e:
            st.error(f"Excel படிக்க முடியவில்லை: {e}")
            return None
    return None

df = load_excel_data()

if df is None or df.empty:
    st.error("❌ 'QA.xlsx' ஃபைல் சரியாக லோட் ஆகவில்லை! ஃபைலை சரிபார்க்கவும்.")
    st.stop()

# --- 2. FSC DETAILS DATA ---
fsc_details = {
    "MURUGAN": "VQ126A0000000196",
    "SOWMIYAN": "VQ126A0000000155",
    "KANNAN": "VQ126A0000000156",
    "ARUN": "VQ126A0000000157",
    "MANIKANDAN": "VQ126A0000000158"
}

# --- 3. BANK LIST DATA ---
bank_options = [
    "HDFC BANK LTD",
    "ICICI BANK LTD",
    "AXIS BANK LTD",
    "STATE BANK OF INDIA",
    "SUNDARAM FINANCE LTD",
    "MAHINDRA FINANCE",
    "CHOLAMANDALAM FINANCE",
    "CANARA BANK",
    "INDIAN OVERSEAS BANK",
    "Other (Type Manually)"
]

# --- 4. VARIANT / MODEL COLUMN FINDER ---
variant_column = None
for col in df.columns:
    col_str = str(col).upper()
    if any(k in col_str for k in ["VARIANT", "VEHICLE", "MODEL", "NAME", "DESCRIPTION"]):
        variant_column = col
        break

if not variant_column:
    variant_column = df.columns[0]

# --- 5. INPUT FIELDS ---
col1, col2 = st.columns(2)

with col1:
    uploaded_docs = st.file_uploader(
        "📄 Upload Documents (Aadhaar Letter / Cards / PDF)", 
        type=["png", "jpg", "jpeg", "pdf"], 
        accept_multiple_files=True
    )
    
    extracted_name = ""
    extracted_address = ""
    
    if uploaded_docs:
        with st.spinner("Processing Document... Reading Text"):
            try:
                import easyocr
                import numpy as np
                
                reader = easyocr.Reader(['en'])
                
                ignore_phrases = [
                    "GOVERNMENT OF INDIA", "GOVERNMENT", "INDIA", "INCOME TAX",
                    "DETAILS AS ON", "UNIQUE IDENTIFICATION", "AUTHORITY", "AADHAAR",
                    "PROOF OF IDENTITY", "CITIZENSHIP", "ENROLLMENT", "HELP@UIDAI",
                    "WWW.UIDAI", "1947", "ISSUE DATE", "YOUR AADHAAR NO", "ENROLLMENT NO"
                ]
                
                raw_text_list = []
                
                for uploaded_doc in uploaded_docs:
                    images_to_process = []
                    
                    if uploaded_doc.name.lower().endswith(".pdf"):
                        import pypdfium2 as pdfium
                        pdf = pdfium.PdfDocument(uploaded_doc.read())
                        for p in pdf:
                            images_to_process.append(p.render(scale=2).to_pil())
                    else:
                        images_to_process.append(Image.open(uploaded_doc))
                    
                    for img in images_to_process:
                        ocr_results = reader.readtext(np.array(img), detail=0)
                        for line in ocr_results:
                            if line.strip():
                                raw_text_list.append(line.strip())
                
                # --- AADHAAR LETTER EXTRACTION (To / S/O Based) ---
                so_index = -1
                for idx, line in enumerate(raw_text_list):
                    l_upper = line.upper()
                    if any(k in l_upper for k in ["S/O", "D/O", "W/O", "C/O"]):
                        so_index = idx
                        break
                
                if so_index != -1:
                    for j in range(so_index - 1, -1, -1):
                        cand = raw_text_list[j].strip()
                        if re.match(r'^[a-zA-Z\s.]+$', cand) and len(cand) >= 3:
                            if not any(ph in cand.upper() for ph in ignore_phrases) and cand.upper() not in ["TO", "TO:"]:
                                extracted_name = cand
                                break
                    
                    addr_lines = []
                    for k in range(so_index, len(raw_text_list)):
                        l_text = raw_text_list[k].strip()
                        l_upper = l_text.upper()
                        
                        if any(ph in l_upper for ph in ignore_phrases) or re.search(r'\d{4}\s?\d{4}\s?\d{4}', l_text) or "MOBILE" in l_upper:
                            break
                        
                        clean_line = re.sub(r'[^a-zA-Z0-9\s,./:-]', '', l_text).strip()
                        if len(clean_line) > 2 and not clean_line.isdigit():
                            addr_lines.append(clean_line)
                    
                    if addr_lines:
                        extracted_address = ", ".join(addr_lines)

                # --- CARD PARSER FALLBACK (DOB / Male Based Name) ---
                if not extracted_name:
                    for i, line in enumerate(raw_text_list):
                        upper_line = line.upper()
                        if any(k in upper_line for k in ["DOB", "DATE OF BIRTH", "MALE", "FEMALE"]):
                            for j in range(i - 1, -1, -1):
                                candidate = raw_text_list[j].strip()
                                if re.match(r'^[a-zA-Z\s.]+$', candidate) and len(candidate) >= 3:
                                    if not any(ph in candidate.upper() for ph in ignore_phrases):
                                        extracted_name = candidate
                                        break
                            if extracted_name:
                                break

                # --- CARD PARSER FALLBACK (Address Block) ---
                if not extracted_address:
                    full_text = " \n ".join(raw_text_list)
                    if "Address" in full_text or "ADDRESS" in full_text:
                        match = re.search(r'(?:Address|ADDRESS)[:\s]*(.*)', full_text, re.DOTALL)
                        if match:
                            addr_block = match.group(1)
                            a_lines = []
                            for line in addr_block.split('\n'):
                                line_str = line.strip()
                                if not re.search(r'(\d{4}\s?\d{4}\s?\d{4}|VID|help@uidai|www|1947|Details as on)', line_str, re.IGNORECASE):
                                    clean_addr = re.sub(r'[^a-zA-Z0-9\s,./:-]', '', line_str).strip()
                                    if len(clean_addr) > 2:
                                        a_lines.append(clean_addr)
                            if a_lines:
                                extracted_address = ", ".join(a_lines[:5])
                
                st.success("Document processed successfully!")
                
            except Exception as e:
                st.error(f"⚠️ Document எக்ஸ்ட்ராக்ட் செய்வதில் பிழை: {e}")

    # Dynamic Value Assignment
    default_name = extracted_name if extracted_name else ""
    default_address = extracted_address if extracted_address else ""

    # Customer Name & Address Inputs
    cust_input = st.text_input("Customer Name", default_name)
    
    if cust_input:
        salutation = get_salutation(cust_input)
        customer_name = f"{salutation} {cust_input.upper()}"
    else:
        customer_name = ""
        
    customer_address = st.text_area("Customer Address", default_address)
    
    # FSC Name Selection
    fsc_name = st.selectbox("Select FSC Name", list(fsc_details.keys()))
    fsc_code = fsc_details[fsc_name]
    st.info(f"Selected FSC Code: **{fsc_code}**")

with col2:
    # 🚗 Vehicle Variant Selection
    vehicle_variants = df[variant_column].dropna().unique()
    selected_variant = st.selectbox(" Select Vehicle Model / Variant", vehicle_variants)
    
    # 🏦 Bank Selection Box
    selected_bank_option = st.selectbox(" Select Hypothecation Bank", bank_options)
    
    if selected_bank_option == "Other (Type Manually)":
        bank_name = st.text_input("Enter Bank Name Manually", "HDFC BANK LTD")
    else:
        bank_name = selected_bank_option

# Selected Row Data
filtered_df = df[df[variant_column] == selected_variant]
model_row = filtered_df.iloc[0] if not filtered_df.empty else pd.Series()

st.divider()

# --- 6. PARTICULARS SELECT & AUTO-FILL PRICES ---
st.subheader("📋 Select Particulars & Auto-Filled Prices")

all_particulars = [
    "Ex Showroom", "Insurance", "KA LIFE TAX", "Life Tax", "Temp", 
    "RSA", "Yellow Paint & Handling", "Fast Tag", "Accessories", "TCS 1 %"
]

selected_particulars = st.multiselect(
    "Quotation-ல் வர வேண்டிய Particulars-ஐ தேர்வு செய்யவும்:",
    options=all_particulars,
    default=["Ex Showroom", "Insurance", "KA LIFE TAX", "TCS 1 %"]
)

items_data = []
total_cost = 0.0

col_p1, col_p2 = st.columns(2)

with col_p1:
    for item in selected_particulars:
        excel_val = 0.0
        
        if not model_row.empty:
            for col in df.columns:
                c_col = str(col).lower().replace(" ", "").replace("_", "")
                c_item = item.lower().replace(" ", "").replace("_", "")
                
                if c_item in c_col or c_col in c_item:
                    try:
                        val = model_row[col]
                        if pd.notna(val):
                            excel_val = float(val)
                    except:
                        excel_val = 0.0
                    break
        
        # Auto-Fill Price
        price = st.number_input(f"Price for {item} (₹)", value=excel_val, step=500.0, key=f"{selected_variant}_{item}")
        
        items_data.append({
            "particulars": item,
            "price": f"{int(price)}" if price.is_integer() else f"{price:,.2f}"
        })
        total_cost += price

formatted_total = f"{int(total_cost)}" if total_cost.is_integer() else f"{total_cost:,.2f}"
st.markdown(f"### 💰 **Total Cost: ₹ {formatted_total}**")

st.divider()

# --- 7. GENERATE WORD & PDF ---
if st.button("🚀 Generate Word & PDF Quotation"):
    template_path = "Quotation_Template.docx"

    if not os.path.exists(template_path):
        st.error("❌ 'Quotation_Template.docx' ஃபால்டரில் இல்லை!")
    else:
        from docx import Document

        doc = Document(template_path)

        # 1. Text Replacements
        replacements = {
            "{{ date }}": pd.Timestamp.now().strftime("%d-%m-%Y"),
            "{{ customer_name }}": customer_name,
            "{{ customer_address }}": customer_address,
            "{{ fsc_name }}": fsc_name,
            "{{ fsc_code }}": fsc_code,
            "{{ bank_name }}": bank_name
        }

        # Replace text in Paragraphs
        for p in doc.paragraphs:
            for key, val in replacements.items():
                if key in p.text:
                    p.text = p.text.replace(key, str(val))

        # Replace & Dynamically Build Table Rows
        for table in doc.tables:
            # Replace Header & Footer Placeholders
            for row in table.rows:
                for cell in row.cells:
                    if "{{ vehicle_variant }}" in cell.text:
                        cell.text = cell.text.replace("{{ vehicle_variant }}", selected_variant)
                    if "{{ total_cost }}" in cell.text:
                        cell.text = cell.text.replace("{{ total_cost }}", formatted_total)

            # Check if table has at least 3 rows
            if len(table.rows) >= 3 and items_data:
                # Target row 1 XML element
                template_tr = table.rows[1]._tr
                last_tr = table.rows[-1]._tr

                # Fill first item in Row 1
                table.rows[1].cells[0].text = items_data[0]["particulars"]
                table.rows[1].cells[1].text = items_data[0]["price"]

                # Clone Row 1 for remaining items (Safe XML insertion)
                for item in items_data[1:]:
                    new_tr = copy.deepcopy(template_tr)
                    last_tr.addprevious(new_tr)

                    # Get cells of new XML row
                    cells = new_tr.xpath('.//w:tc')
                    if len(cells) >= 2:
                        p0 = cells[0].xpath('.//w:p')[0]
                        p0.text = item["particulars"]
                        p1 = cells[1].xpath('.//w:p')[0]
                        p1.text = item["price"]

        word_filename = f"Quotation_{customer_name}.docx"
        pdf_filename = f"Quotation_{customer_name}.pdf"

        doc.save(word_filename)

        # PDF Conversion
        pdf_ready = False
        try:
            from docx2pdf import convert
            convert(word_filename, pdf_filename)
            pdf_ready = True
        except:
            try:
                import comtypes.client
                word = comtypes.client.CreateObject("Word.Application")
                word.Visible = False
                doc_path = os.path.abspath(word_filename)
                pdf_path = os.path.abspath(pdf_filename)

                doc_pdf = word.Documents.Open(doc_path)
                doc_pdf.SaveAs(pdf_path, FileFormat=17)
                doc_pdf.Close()
                word.Quit()
                pdf_ready = True
            except Exception as e:
                st.warning(f"⚠️ PDF உருவாக்க முடியவில்லை: {e}")

        st.success("🎉 Quotation வெற்றிகரமாக உருவாக்கப்பட்டது!")

        c1, c2 = st.columns(2)
        with c1:
            with open(word_filename, "rb") as f_word:
                st.download_button(
                    label="📥 Download WORD File",
                    data=f_word,
                    file_name=word_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        with c2:
            if pdf_ready and os.path.exists(pdf_filename):
                with open(pdf_filename, "rb") as f_pdf:
                    st.download_button(
                        label="📄 Download PDF File",
                        data=f_pdf,
                        file_name=pdf_filename,
                        mime="application/pdf"
                    )
